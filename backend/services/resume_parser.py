"""
Resume Parser Service
Extracts text from PDF and DOCX files
"""

import pdfplumber
from docx import Document
import re
from typing import Optional


class ResumeParser:
    """
    Service for extracting text from resume files (PDF/DOCX)
    """
    
    def extract_text(self, file_path: str, file_extension: str) -> str:
        """
        Extract text from a resume file based on its extension
        
        Args:
            file_path: Path to the temporary file
            file_extension: File extension (.pdf or .docx)
            
        Returns:
            Extracted text content
        """
        if file_extension.lower() == '.pdf':
            return self._extract_from_pdf(file_path)
        elif file_extension.lower() == '.docx':
            return self._extract_from_docx(file_path)
        else:
            raise ValueError(f"Unsupported file type: {file_extension}")
    
    def _extract_from_pdf(self, file_path: str) -> str:
        """
        Extract text from PDF using pdfplumber
        """
        text_content = []
        
        try:
            with pdfplumber.open(file_path) as pdf:
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text_content.append(page_text)
            
            return '\n'.join(text_content)
            
        except Exception as e:
            raise Exception(f"Failed to extract text from PDF: {str(e)}")
    
    def _extract_from_docx(self, file_path: str) -> str:
        """
        Extract text from DOCX using python-docx
        """
        text_content = []
        
        try:
            doc = Document(file_path)
            
            # Extract paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_content.append(paragraph.text)
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        text_content.append(' | '.join(row_text))
            
            return '\n'.join(text_content)
            
        except Exception as e:
            raise Exception(f"Failed to extract text from DOCX: {str(e)}")
    
    def extract_contact_info(self, text: str) -> dict:
        """
        Extract contact information from resume text
        """
        contact_info = {
            'email': None,
            'phone': None,
            'linkedin': None,
            'github': None
        }
        
        # Email pattern
        email_pattern = r'[\w\.-]+@[\w\.-]+\.\w+'
        email_match = re.search(email_pattern, text)
        if email_match:
            contact_info['email'] = email_match.group()
        
        # Phone pattern (various formats)
        phone_pattern = r'[\+]?[(]?[0-9]{1,3}[)]?[-\s\.]?[(]?[0-9]{1,4}[)]?[-\s\.]?[0-9]{1,4}[-\s\.]?[0-9]{1,9}'
        phone_match = re.search(phone_pattern, text)
        if phone_match:
            contact_info['phone'] = phone_match.group()
        
        # LinkedIn pattern
        linkedin_pattern = r'linkedin\.com/in/[\w-]+'
        linkedin_match = re.search(linkedin_pattern, text, re.IGNORECASE)
        if linkedin_match:
            contact_info['linkedin'] = linkedin_match.group()
        
        # GitHub pattern
        github_pattern = r'github\.com/[\w-]+'
        github_match = re.search(github_pattern, text, re.IGNORECASE)
        if github_match:
            contact_info['github'] = github_match.group()
        
        return contact_info
