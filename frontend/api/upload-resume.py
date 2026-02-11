"""
Vercel Serverless Function: Upload Resume
"""

from http.server import BaseHTTPRequestHandler
import json
import tempfile
import os
import sys

# Add parent directory for imports
sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

try:
    import pdfplumber
    from docx import Document
except ImportError:
    pdfplumber = None
    Document = None

MAX_FILE_SIZE = 5 * 1024 * 1024  # 5MB
ALLOWED_EXTENSIONS = {'.pdf', '.docx'}


def extract_from_pdf(file_path: str) -> str:
    """Extract text from PDF"""
    text_content = []
    with pdfplumber.open(file_path) as pdf:
        for page in pdf.pages:
            page_text = page.extract_text()
            if page_text:
                text_content.append(page_text)
    return '\n'.join(text_content)


def extract_from_docx(file_path: str) -> str:
    """Extract text from DOCX"""
    text_content = []
    doc = Document(file_path)
    for paragraph in doc.paragraphs:
        if paragraph.text.strip():
            text_content.append(paragraph.text)
    for table in doc.tables:
        for row in table.rows:
            row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if row_text:
                text_content.append(' | '.join(row_text))
    return '\n'.join(text_content)


def clean_text(text: str) -> str:
    """Clean extracted text"""
    import re
    replacements = {
        '\u2019': "'", '\u2018': "'", '\u201c': '"', '\u201d': '"',
        '\u2013': '-', '\u2014': '-', '\u2022': '*', '\u2026': '...',
        '\xa0': ' ', '\t': ' ',
    }
    for old, new in replacements.items():
        text = text.replace(old, new)
    text = re.sub(r'[^\w\s\-\.\,\;\:\!\?\(\)\[\]\{\}\/\&\+\#\@\'\"\|]', ' ', text)
    text = re.sub(r'\s+', ' ', text)
    return text.strip()


class handler(BaseHTTPRequestHandler):
    def do_POST(self):
        try:
            content_length = int(self.headers.get('Content-Length', 0))
            content_type = self.headers.get('Content-Type', '')
            
            if content_length > MAX_FILE_SIZE:
                self.send_error_response(400, "File size exceeds 5MB limit")
                return
            
            # Parse multipart form data
            body = self.rfile.read(content_length)
            
            # Extract boundary from content type
            boundary = None
            for part in content_type.split(';'):
                if 'boundary' in part:
                    boundary = part.split('=')[1].strip()
                    break
            
            if not boundary:
                self.send_error_response(400, "Invalid multipart form data")
                return
            
            # Parse the multipart data
            parts = body.split(f'--{boundary}'.encode())
            file_content = None
            filename = None
            
            for part in parts:
                if b'filename=' in part:
                    # Extract filename
                    header_end = part.find(b'\r\n\r\n')
                    if header_end != -1:
                        header = part[:header_end].decode('utf-8', errors='ignore')
                        for line in header.split('\r\n'):
                            if 'filename=' in line:
                                start = line.find('filename="') + 10
                                end = line.find('"', start)
                                filename = line[start:end]
                        file_content = part[header_end + 4:].rstrip(b'\r\n--')
                        break
            
            if not file_content or not filename:
                self.send_error_response(400, "No file uploaded")
                return
            
            # Validate file extension
            file_ext = os.path.splitext(filename)[1].lower()
            if file_ext not in ALLOWED_EXTENSIONS:
                self.send_error_response(400, f"Invalid file type. Allowed: {', '.join(ALLOWED_EXTENSIONS)}")
                return
            
            # Save to temp file and extract text
            with tempfile.NamedTemporaryFile(delete=False, suffix=file_ext) as temp_file:
                temp_file.write(file_content)
                temp_path = temp_file.name
            
            try:
                if file_ext == '.pdf':
                    extracted_text = extract_from_pdf(temp_path)
                else:
                    extracted_text = extract_from_docx(temp_path)
                
                cleaned_text = clean_text(extracted_text)
                
                if len(cleaned_text) < 50:
                    self.send_error_response(400, "Could not extract sufficient text from resume")
                    return
                
                self.send_json_response({
                    "success": True,
                    "filename": filename,
                    "extracted_text": cleaned_text,
                    "character_count": len(cleaned_text),
                    "word_count": len(cleaned_text.split())
                })
            finally:
                os.unlink(temp_path)
                
        except Exception as e:
            self.send_error_response(500, str(e))
    
    def send_json_response(self, data, status=200):
        self.send_response(status)
        self.send_header('Content-Type', 'application/json')
        self.send_header('Access-Control-Allow-Origin', '*')
        self.send_header('Access-Control-Allow-Methods', 'POST, OPTIONS')
        self.send_header('Access-Control-Allow-Headers', 'Content-Type')
        self.end_headers()
        self.wfile.write(json.dumps(data).encode())
    
    def send_error_response(self, status, message):
        self.send_json_response({"detail": message}, status)
    
    def do_OPTIONS(self):
        self.send_response(200)
        self.send_header('Access-Control-Allow-Origin', '*')
        self.send_header('Access-Control-Allow-Methods', 'POST, OPTIONS')
        self.send_header('Access-Control-Allow-Headers', 'Content-Type')
        self.end_headers()
